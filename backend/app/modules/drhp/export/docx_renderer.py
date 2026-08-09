"""Render assembled DRHP export documents to DOCX (python-docx)."""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from app.modules.drhp.ast.schemas import DrhpBlockAST, DrhpSectionAST
from app.modules.drhp.export.document import COVER_CHAPTER_KEY, DRHPExportDocument, ExportChapter
from app.modules.drhp.export.formatters import chapter_title_duplicates_section, normalize_heading_text
from app.modules.drhp.export.publication_theme import DRAFT_FOOTER_NOTICE, DRAFT_HEADER, PUBLICATION_THEME

BULLET_CHAR = "\u2022"


def _set_run_font(run, *, size_pt: float, bold: bool = False, italic: bool = False) -> None:  # noqa: ANN001
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic


def _set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 4) -> None:  # noqa: ANN001
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


class DRHPDocxRenderer:
    def __init__(self) -> None:
        self.theme = PUBLICATION_THEME
        self._is_cover = False
        self._is_risk_chapter = False

    def render(self, document: DRHPExportDocument) -> bytes:
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(22)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(18)
        section.right_margin = Mm(18)
        self._configure_header_footer(section, document)
        self._build_core_properties(doc, document)

        cover = document.cover_chapter()
        if cover is not None:
            self._is_cover = True
            self._build_chapter(doc, cover)
            doc.add_page_break()
            self._is_cover = False

        self._build_toc(doc, document)
        doc.add_page_break()

        for index, chapter in enumerate(document.body_chapters()):
            if index > 0:
                doc.add_page_break()
            self._is_risk_chapter = chapter.chapter_key == "risk-factors"
            self._build_chapter(doc, chapter)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _configure_header_footer(self, section, document: DRHPExportDocument) -> None:  # noqa: ANN001
        header = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        header.text = DRAFT_HEADER.upper()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header.runs:
            _set_run_font(header.runs[0], size_pt=self.theme.header_size_pt, bold=False)
        if document.issuer_name:
            issuer_para = section.header.add_paragraph(document.issuer_name)
            issuer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if issuer_para.runs:
                _set_run_font(issuer_para.runs[0], size_pt=self.theme.header_size_pt)

        footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer.text = DRAFT_FOOTER_NOTICE
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer.runs:
            _set_run_font(footer.runs[0], size_pt=self.theme.footer_size_pt, italic=True)

        page_field = section.footer.add_paragraph()
        page_field.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = page_field.add_run("Page ")
        _set_run_font(run, size_pt=self.theme.footer_size_pt)
        self._add_page_number_field(page_field)

    def _add_page_number_field(self, paragraph) -> None:  # noqa: ANN001
        run = paragraph.add_run()
        fld_char_begin = self._field_element("begin")
        instr = self._field_element("instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_char_sep = self._field_element("separate")
        fld_char_end = self._field_element("end")
        run._r.extend([fld_char_begin, instr, fld_char_sep, fld_char_end])

    @staticmethod
    def _field_element(field_type: str):
        from docx.oxml import OxmlElement

        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), field_type)
        return fld

    def _build_core_properties(self, doc: Document, document: DRHPExportDocument) -> None:
        props = doc.core_properties
        props.title = document.document_title
        props.subject = "Draft Red Herring Prospectus"
        if document.issuer_name:
            props.author = document.issuer_name
        if document.generated_at:
            props.created = document.generated_at

    def _build_toc(self, doc: Document, document: DRHPExportDocument) -> None:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run("TABLE OF CONTENTS")
        _set_run_font(run, size_pt=self.theme.toc_title_size_pt, bold=True)
        _set_paragraph_spacing(title, after=8)
        if document.partial_label:
            para = doc.add_paragraph(f"({document.partial_label})")
            if para.runs:
                _set_run_font(para.runs[0], size_pt=self.theme.body_size_pt, italic=True)
        for entry in document.table_of_contents:
            para = doc.add_paragraph(f"{entry.display_number}. {entry.title}")
            para.paragraph_format.left_indent = Mm(5)
            _set_paragraph_spacing(para, after=1)
            if para.runs:
                _set_run_font(para.runs[0], size_pt=self.theme.body_size_pt)
        notice = doc.add_paragraph(DRAFT_FOOTER_NOTICE)
        if notice.runs:
            _set_run_font(notice.runs[0], size_pt=self.theme.legal_notice_size_pt, italic=True)

    def _should_emit_chapter_title(self, chapter: ExportChapter) -> bool:
        if not chapter.available or chapter.chapter_ast is None:
            return True
        if chapter.chapter_key == COVER_CHAPTER_KEY:
            return False
        first_section = chapter.chapter_ast.sections[0] if chapter.chapter_ast.sections else None
        if first_section and first_section.heading:
            if chapter_title_duplicates_section(chapter.title, normalize_heading_text(first_section.heading)):
                return False
        return True

    def _build_chapter(self, doc: Document, chapter: ExportChapter) -> None:
        if self._should_emit_chapter_title(chapter):
            heading = doc.add_heading(chapter.title, level=1)
            if heading.runs:
                _set_run_font(heading.runs[0], size_pt=self.theme.chapter_title_size_pt, bold=True)
            _set_paragraph_spacing(heading, before=0, after=6)
        if not chapter.available or chapter.chapter_ast is None:
            reason = chapter.unavailable_reason or "This chapter is unavailable in this draft."
            para = doc.add_paragraph(reason)
            if para.runs:
                _set_run_font(para.runs[0], size_pt=self.theme.body_size_pt, italic=True)
            return

        for section in chapter.chapter_ast.sections:
            self._build_section(doc, section)

    def _build_section(self, doc: Document, section: DrhpSectionAST) -> None:
        if section.heading:
            heading = doc.add_heading(section.heading, level=2)
            if heading.runs:
                _set_run_font(heading.runs[0], size_pt=self.theme.section_heading_size_pt, bold=True)
            _set_paragraph_spacing(heading, before=6, after=4)
            heading.paragraph_format.keep_with_next = True
        for block in section.blocks:
            self._build_block(doc, block)

    def _build_block(self, doc: Document, block: DrhpBlockAST) -> None:
        if block.kind == "page_break":
            doc.add_page_break()
            return

        content = block.content or {}
        if block.kind == "heading":
            level = int(content.get("level") or 3)
            size = self.theme.risk_title_size_pt if self._is_risk_chapter else self.theme.subsection_heading_size_pt
            heading = doc.add_heading(str(content.get("text") or ""), level=min(max(level, 1), 3))
            if heading.runs:
                _set_run_font(heading.runs[0], size_pt=size, bold=True)
            heading.paragraph_format.keep_with_next = True
            return

        if block.kind in {"paragraph", "placeholder", "cross_reference"}:
            text = str(content.get("text") or "")
            para = doc.add_paragraph()
            if self._is_cover:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if text.upper() == DRAFT_HEADER:
                    run = para.add_run(text)
                    _set_run_font(run, size_pt=self.theme.cover_label_size_pt, bold=True)
                    return
            run = para.add_run(text)
            _set_run_font(run, size_pt=self.theme.body_size_pt)
            para.paragraph_format.line_spacing = 1.15
            _set_paragraph_spacing(para, after=4)
            return

        if block.kind == "legal_notice":
            para = doc.add_paragraph(str(content.get("text") or ""))
            if para.runs:
                _set_run_font(para.runs[0], size_pt=self.theme.legal_notice_size_pt, italic=True)
            _set_paragraph_spacing(para, after=5)
            return

        if block.kind == "bullet_list":
            for item in content.get("items") or []:
                text = str(item).strip()
                if not text:
                    continue
                para = doc.add_paragraph(f"{BULLET_CHAR} {text}")
                para.paragraph_format.left_indent = Mm(5)
                if para.runs:
                    _set_run_font(para.runs[0], size_pt=self.theme.body_size_pt)
                _set_paragraph_spacing(para, after=2)
            return

        if block.kind == "numbered_list":
            for index, item in enumerate(content.get("items") or [], start=1):
                text = str(item).strip()
                if not text:
                    continue
                para = doc.add_paragraph(f"{index}. {text}")
                para.paragraph_format.left_indent = Mm(5)
                if para.runs:
                    _set_run_font(para.runs[0], size_pt=self.theme.body_size_pt)
                _set_paragraph_spacing(para, after=2)
            return

        if block.kind in {"table", "key_value_table"}:
            self._build_table(doc, content)
            return

        para = doc.add_paragraph(str(content.get("text") or ""))
        if para.runs:
            _set_run_font(para.runs[0], size_pt=self.theme.body_size_pt)

    def _build_table(self, doc: Document, content: dict[str, Any]) -> None:
        headers = [str(item) for item in content.get("headers") or []]
        rows = content.get("rows") or []
        caption = str(content.get("caption") or "").strip()
        notes = content.get("notes") or []
        alignments = content.get("columnAlignments") or []
        unit = str(content.get("unit") or "").strip()

        if caption:
            cap = doc.add_paragraph(caption)
            if cap.runs:
                _set_run_font(cap.runs[0], size_pt=self.theme.table_caption_size_pt, bold=True)
            cap.paragraph_format.keep_with_next = True
        if unit and unit.casefold() not in caption.casefold():
            unit_para = doc.add_paragraph(f"(₹ in {unit}, unless otherwise stated)")
            if unit_para.runs:
                _set_run_font(unit_para.runs[0], size_pt=self.theme.table_note_size_pt)

        body_rows = len(rows) + (1 if headers else 0)
        if body_rows == 0:
            return
        col_count = len(headers) if headers else max((len(row) for row in rows if isinstance(row, list)), default=0)
        if col_count == 0:
            return

        table = doc.add_table(rows=body_rows, cols=col_count)
        table.style = "Table Grid"
        row_index = 0
        if headers:
            for col_index, header in enumerate(headers):
                cell = table.rows[0].cells[col_index]
                cell.text = header
                if cell.paragraphs and cell.paragraphs[0].runs:
                    _set_run_font(cell.paragraphs[0].runs[0], size_pt=self.theme.table_size_pt, bold=True)
            row_index = 1

        for row in rows:
            if not isinstance(row, list):
                continue
            if row_index >= len(table.rows):
                break
            for col_index, value in enumerate(row[:col_count]):
                cell = table.rows[row_index].cells[col_index]
                cell.text = str(value)
                if cell.paragraphs and cell.paragraphs[0].runs:
                    _set_run_font(cell.paragraphs[0].runs[0], size_pt=self.theme.table_size_pt)
                alignment = alignments[col_index] if col_index < len(alignments) else "left"
                if alignment == "right" and cell.paragraphs:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_index += 1

        if notes:
            notes_para = doc.add_paragraph("Notes:")
            if notes_para.runs:
                _set_run_font(notes_para.runs[0], size_pt=self.theme.table_note_size_pt, bold=True)
            for index, note in enumerate(notes, start=1):
                note_para = doc.add_paragraph(f"{index}. {note}")
                if note_para.runs:
                    _set_run_font(note_para.runs[0], size_pt=self.theme.table_note_size_pt)


def render_docx(document: DRHPExportDocument) -> bytes:
    return DRHPDocxRenderer().render(document)
